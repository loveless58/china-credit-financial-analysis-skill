#!/usr/bin/env python3
"""Insert or replace a financial-analysis update in a DOCX file."""

from __future__ import annotations

import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - environment guard
    raise SystemExit("python-docx is required: pip install python-docx") from exc

SECTION_KEYWORDS = ("财务分析", "财务状况", "资产负债", "盈利能力", "现金流")


def find_anchor(document: Document) -> int | None:
    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if any(keyword in text for keyword in SECTION_KEYWORDS):
            return idx
    return None


def add_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.add_run(text)
    return inserted


def find_paragraph_index(document: Document, anchor: str, start: int = 0) -> int:
    for index, paragraph in enumerate(document.paragraphs[start:], start):
        if anchor in paragraph.text:
            return index
    raise ValueError(f"paragraph anchor not found: {anchor}")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_run(run, font_name: str, font_size: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), font_name
    )


def locate_financial_section_anchors(
    document: Document,
    section_start: str,
    analysis_anchor: str,
    section_end: str,
) -> tuple[int, int, int]:
    section_index = find_paragraph_index(document, section_start)
    end_index = find_paragraph_index(document, section_end, section_index + 1)
    analysis_matches = [
        index
        for index in range(section_index + 1, end_index)
        if analysis_anchor in document.paragraphs[index].text
    ]
    if len(analysis_matches) != 1:
        raise ValueError(
            "expected exactly one analysis anchor inside nearest financial section; "
            f"found {len(analysis_matches)}; analysis_anchor={analysis_anchor!r}"
        )
    analysis_index = analysis_matches[0]
    return section_index, analysis_index, end_index


def _insert_formatted_lines(
    anchor,
    lines: list[str],
    shading: str,
    body_font: str,
    body_size: float,
) -> None:
    current = anchor
    for line in lines:
        if not line.strip():
            continue
        current = add_paragraph_after(current, line)
        run = current.runs[-1]
        set_paragraph_shading(current, shading)
        format_run(run, body_font, body_size)


def replace_analysis_body(
    document: Document,
    section_start: str,
    analysis_anchor: str,
    section_end: str,
    analysis_markdown: str,
    shading: str,
    body_font: str = "仿宋_GB2312",
    body_size: float = 14,
) -> str:
    _, analysis_index, end_index = locate_financial_section_anchors(
        document, section_start, analysis_anchor, section_end
    )
    paragraphs = document.paragraphs
    for paragraph in reversed(paragraphs[analysis_index + 1 : end_index]):
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
    _insert_formatted_lines(
        paragraphs[analysis_index],
        analysis_markdown.splitlines(),
        shading,
        body_font,
        body_size,
    )
    return f"after analysis anchor paragraph {analysis_index + 1} (replace)"


def insert_analysis_body(
    document: Document,
    section_start: str,
    analysis_anchor: str,
    section_end: str,
    analysis_markdown: str,
    shading: str,
    body_font: str = "仿宋_GB2312",
    body_size: float = 14,
) -> str:
    _, analysis_index, _ = locate_financial_section_anchors(
        document, section_start, analysis_anchor, section_end
    )
    lines = ["【Codex财务分析更新建议】", *analysis_markdown.splitlines()]
    _insert_formatted_lines(
        document.paragraphs[analysis_index],
        lines,
        shading,
        body_font,
        body_size,
    )
    return f"after analysis anchor paragraph {analysis_index + 1} (insert)"


def insert_block(document: Document, anchor_idx: int | None, update_text: str) -> str:
    marker = "【Codex财务分析更新建议】"
    pending_marker = "【待核验事项】"
    lines = [marker, *update_text.splitlines()]
    if pending_marker not in update_text:
        lines.extend(["", pending_marker, "详见待核验清单。"])

    if anchor_idx is None:
        raise ValueError("financial-analysis section not found; insertion is blocked")

    anchor = document.paragraphs[anchor_idx]
    current = anchor
    for line in lines:
        current = add_paragraph_after(current, line)
    return f"after paragraph {anchor_idx + 1}: {anchor.text[:80]}"


def replace_section(document: Document, anchor_idx: int | None, update_text: str) -> str:
    if anchor_idx is None:
        raise SystemExit("financial-analysis section not found; replacement is blocked")
    paragraphs = document.paragraphs
    start = anchor_idx
    end = len(paragraphs)
    for idx in range(anchor_idx + 1, len(paragraphs)):
        style_name = getattr(paragraphs[idx].style, "name", "")
        text = paragraphs[idx].text.strip()
        if idx > start + 1 and (style_name.startswith("Heading") or text.startswith(("一、", "二、", "三、", "四、", "五、", "六、"))):
            end = idx
            break

    for idx in range(end - 1, start, -1):
        element = paragraphs[idx]._element
        element.getparent().remove(element)

    paragraphs[start].text = ""
    current = paragraphs[start]
    for line in update_text.splitlines():
        if not current.text:
            current.add_run(line)
        else:
            current = add_paragraph_after(current, line)
    return f"replaced paragraphs {start + 1}-{end}"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("update_markdown", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--mode", choices=("insert", "replace"), default="insert")
    parser.add_argument("--confirm-replace", action="store_true")
    parser.add_argument("--section-start", default="财务分析")
    parser.add_argument("--analysis-anchor", default="合并财务情况分析")
    parser.add_argument("--section-end", default="行业分析")
    parser.add_argument("--change-shading", default="FFF2CC")
    args = parser.parse_args()

    if args.docx.suffix.lower() != ".docx":
        raise SystemExit("only .docx files are supported")
    if args.out.exists() and args.out.resolve() == args.docx.resolve():
        raise SystemExit("refusing to overwrite original path; write to a new file or back up explicitly")
    if args.mode == "replace" and not args.confirm_replace:
        raise SystemExit("replacement requires --confirm-replace")

    document = Document(str(args.docx))
    update_text = args.update_markdown.read_text(encoding="utf-8-sig").strip()
    try:
        if args.mode == "insert":
            location = insert_analysis_body(
                document,
                args.section_start,
                args.analysis_anchor,
                args.section_end,
                update_text,
                args.change_shading,
            )
        else:
            location = replace_analysis_body(
                document,
                args.section_start,
                args.analysis_anchor,
                args.section_end,
                update_text,
                args.change_shading,
            )
    except ValueError as error:
        raise SystemExit(str(error)) from error

    args.out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.out))
    print(location)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
