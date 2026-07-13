#!/usr/bin/env python3
"""Clone a DOCX financial summary table and replace cell text only.

Use this for asset-liability summary tables whose font and table formatting
must match the original credit-review report template.
"""

from __future__ import annotations

import argparse
import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


DEFAULT_ANCHORS = ("资产负债简表", "资产总计", "负债合计", "利润及利润分配表", "现金流量简表")


def table_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def find_table(document: Document, anchors: list[str]) -> tuple[int, Table]:
    for index, table in enumerate(document.tables):
        text = table_text(table)
        if any(anchor in text for anchor in anchors):
            return index, table
    raise SystemExit(f"financial template table not found; anchors={anchors}")


def _document_body_items(document: Document) -> list[Paragraph | Table]:
    items: list[Paragraph | Table] = []
    for child in document._body._body.iterchildren():
        if child.tag == qn("w:p"):
            items.append(Paragraph(child, document._body))
        elif child.tag == qn("w:tbl"):
            items.append(Table(child, document._body))
    return items


def tables_in_section(
    document: Document,
    section_start: str,
    section_end: str,
) -> list[tuple[int, Table]]:
    """Return top-level tables bounded by section paragraphs in body order."""
    items = _document_body_items(document)
    start_index = next(
        (
            index
            for index, item in enumerate(items)
            if isinstance(item, Paragraph) and section_start in item.text
        ),
        None,
    )
    if start_index is None:
        raise ValueError(f"section start not found: {section_start}")
    end_index = next(
        (
            index
            for index, item in enumerate(items[start_index + 1 :], start_index + 1)
            if isinstance(item, Paragraph) and section_end in item.text
        ),
        None,
    )
    if end_index is None:
        raise ValueError(f"section end not found after start: {section_end}")

    document_table_indexes = {
        id(table._tbl): index for index, table in enumerate(document.tables)
    }
    scoped: list[tuple[int, Table]] = []
    for item in items[start_index + 1 : end_index]:
        if not isinstance(item, Table):
            continue
        table_index = document_table_indexes.get(id(item._tbl))
        if table_index is None:
            raise ValueError("scoped table is missing from document table index")
        scoped.append((table_index, item))
    return scoped


def find_table_in_section(
    document: Document,
    anchors: list[str],
    section_start: str,
    section_end: str,
    preceding_anchor: str | None = None,
) -> tuple[int, Table]:
    """Locate exactly one anchored table inside the requested body section."""
    anchored_matches = [
        (index, table)
        for index, table in tables_in_section(document, section_start, section_end)
        if any(anchor in table_text(table) for anchor in anchors)
    ]
    if len(anchored_matches) == 1:
        return anchored_matches[0]

    matches = anchored_matches
    if preceding_anchor is not None:
        preceding_text: dict[int, str] = {}
        last_paragraph = ""
        for item in _document_body_items(document):
            if isinstance(item, Paragraph):
                if item.text.strip():
                    last_paragraph = item.text.strip()
            else:
                preceding_text[id(item._tbl)] = last_paragraph
        contextual_matches = [
            (index, table)
            for index, table in anchored_matches
            if preceding_anchor in preceding_text.get(id(table._tbl), "")
        ]
        if len(contextual_matches) == 1:
            return contextual_matches[0]
        if contextual_matches:
            matches = contextual_matches
    if len(matches) != 1:
        raise ValueError(
            "expected exactly one financial template table inside section "
            f"{section_start!r}..{section_end!r}; found {len(matches)}; "
            f"anchors={anchors}; preceding_anchor={preceding_anchor!r}"
        )
    return matches[0]


def text_nodes(cell) -> list:
    return list(cell._tc.iter(qn("w:t")))


def set_cell_text_preserving_format(cell, text: str) -> None:
    nodes = text_nodes(cell)
    if nodes:
        nodes[0].text = text
        for node in nodes[1:]:
            node.text = ""
        return
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    run._element.get_or_add_rPr()


def fill_table(table: Table, rows: list[list[str]]) -> None:
    if len(rows) > len(table.rows):
        raise SystemExit(f"row count exceeds template: data={len(rows)} template={len(table.rows)}")
    for row_index, row in enumerate(rows):
        if len(row) > len(table.rows[row_index].cells):
            raise SystemExit(
                "column count exceeds template at row "
                f"{row_index}: data={len(row)} template={len(table.rows[row_index].cells)}"
            )
        for col_index, value in enumerate(row):
            set_cell_text_preserving_format(table.cell(row_index, col_index), str(value))


def clone_table(template: Table, mode: str) -> Table:
    cloned_tbl = deepcopy(template._tbl)
    if mode == "replace-template":
        template._tbl.addprevious(cloned_tbl)
        parent = template._tbl.getparent()
        parent.remove(template._tbl)
    else:
        template._tbl.addnext(cloned_tbl)
        cloned_tbl.addnext(OxmlElement("w:p"))
    return Table(cloned_tbl, template._parent)


def load_rows(rows_json: Path, key: str) -> list[list[str]]:
    rows_data = json.loads(rows_json.read_text(encoding="utf-8-sig"))
    rows = rows_data.get(key)
    if not isinstance(rows, list) or not all(isinstance(row, list) for row in rows):
        raise SystemExit(f"rows_json key must contain a 2D row array: {key}")
    return [[str(cell) for cell in row] for row in rows]


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("rows_json", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--table-key", required=True)
    parser.add_argument("--anchor", action="append", help="Anchor text used to locate the template table")
    parser.add_argument("--mode", choices=("append-after-template", "replace-template"), default="append-after-template")
    args = parser.parse_args()

    if args.docx.suffix.lower() != ".docx":
        raise SystemExit("only .docx files are supported")
    if args.out.resolve() == args.docx.resolve():
        raise SystemExit("refusing to overwrite input DOCX")

    rows = load_rows(args.rows_json, args.table_key)
    document = Document(str(args.docx))
    anchors = args.anchor or list(DEFAULT_ANCHORS)
    table_index, template = find_table(document, anchors)
    target = clone_table(template, args.mode)
    fill_table(target, rows)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.out))
    print(json.dumps({"output": str(args.out), "template_table_index": table_index, "mode": args.mode}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
