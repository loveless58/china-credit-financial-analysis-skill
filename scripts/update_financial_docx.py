#!/usr/bin/env python3
"""Safely update a credit-review financial-analysis DOCX from a bundle."""

from __future__ import annotations

import argparse
import json
import tempfile
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if __package__:
    from .audit_report_numbers import audit_text
    from .backup_docx import backup_path, create_backup
    from .export_change_log import write_change_log
    from .insert_financial_analysis import (
        insert_analysis_body,
        locate_financial_section_anchors,
        replace_analysis_body,
    )
    from .preserve_financial_table_format import fill_table, find_table_in_section
    from .validate_financial_analysis_bundle import validate_bundle
    from .validate_financial_docx import validate_financial_docx
else:
    from audit_report_numbers import audit_text
    from backup_docx import backup_path, create_backup
    from export_change_log import write_change_log
    from insert_financial_analysis import (
        insert_analysis_body,
        locate_financial_section_anchors,
        replace_analysis_body,
    )
    from preserve_financial_table_format import fill_table, find_table_in_section
    from validate_financial_analysis_bundle import validate_bundle
    from validate_financial_docx import validate_financial_docx


SKILL_ROOT = Path(__file__).resolve().parent.parent


class UpdateBlocked(RuntimeError):
    """Raised before a writable DOCX can be safely produced."""

    def __init__(self, message: str, details: list[Any] | None = None) -> None:
        super().__init__(message)
        self.details = details or []


class UpdateFailed(RuntimeError):
    """Raised when the produced DOCX fails structural validation."""

    def __init__(self, message: str, details: list[Any] | None = None) -> None:
        super().__init__(message)
        self.details = details or []


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8-sig") as file_handle:
        return json.load(file_handle)


def write_json(path: Path, payload: object) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return path


def reject_out_dir_inside_skill_repo(out_dir: Path) -> None:
    resolved = out_dir.resolve()
    if resolved == SKILL_ROOT or SKILL_ROOT in resolved.parents:
        raise UpdateBlocked(
            "output directory must be outside the Skill repository",
            [str(resolved)],
        )


def write_pending_markdown(
    path: Path,
    pending: list[dict[str, Any]],
    runtime_failures: list[str] | None = None,
) -> Path:
    lines = ["# 待核验清单", ""]
    if pending:
        lines.extend(
            f"- [{item['status']}] {item['issue']}" for item in pending
        )
    else:
        lines.append("- 无")
    if runtime_failures:
        lines.extend(["", "## 运行失败原因", ""])
        lines.extend(f"- {failure}" for failure in runtime_failures)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def _same_file_entity(first: Path, second: Path) -> bool:
    if not first.exists() or not second.exists():
        return False
    try:
        return first.samefile(second)
    except OSError:
        return False


def _ensure_distinct_document_paths(
    source: Path,
    output: Path,
    backup: Path,
) -> None:
    paths = {"source": source.resolve(), "output": output.resolve(), "backup": backup.resolve()}
    names = list(paths)
    for index, first_name in enumerate(names):
        for second_name in names[index + 1 :]:
            first = paths[first_name]
            second = paths[second_name]
            if first == second or _same_file_entity(first, second):
                raise UpdateBlocked(
                    f"{first_name} and {second_name} refer to the same file entity",
                    [str(first), str(second)],
                )


def _safe_output_path(out_dir: Path, output_filename: str, source: Path) -> Path:
    filename = Path(output_filename)
    if filename.is_absolute() or filename.name != output_filename:
        raise UpdateBlocked("output_filename must be a DOCX basename", [output_filename])
    output = (out_dir / filename).resolve()
    if output.exists():
        if _same_file_entity(source, output):
            raise UpdateBlocked(
                "output path already exists and is the same file entity as source",
                [str(output), str(source.resolve())],
            )
        raise UpdateBlocked("output path already exists", [str(output)])
    if output == source.resolve():
        raise UpdateBlocked(
            "output path resolves to source",
            [str(output), str(source.resolve())],
        )
    return output


def _staging_output_path(output: Path) -> Path:
    with tempfile.NamedTemporaryFile(
        dir=output.parent,
        prefix=f".{output.name}.staging-",
        suffix=".docx",
        delete=False,
    ) as file_handle:
        return Path(file_handle.name)


def _validate_table_dimensions(table, rows: list[list[str]]) -> None:
    template_row_count = len(table.rows)
    if len(rows) != template_row_count:
        raise UpdateBlocked(
            "table dimensions do not match template",
            [f"rows: data={len(rows)} template={template_row_count}"],
        )
    for row_index, row in enumerate(rows):
        template_col_count = len(table.rows[row_index].cells)
        if len(row) != template_col_count:
            raise UpdateBlocked(
                "table dimensions do not match template",
                [
                    f"row {row_index} columns: "
                    f"data={len(row)} template={template_col_count}"
                ],
            )


def phase1_asset_liability_rows(
    plan: dict[str, Any],
) -> list[list[str]] | None:
    if plan.get("preserve_asset_liability_table") is not True:
        raise ValueError("preserve_asset_liability_table must equal true")
    table_rows = plan.get("table_rows")
    if not isinstance(table_rows, dict):
        raise ValueError("table_rows must be an object")
    unsupported = set(table_rows) - {"asset_liability"}
    if unsupported:
        raise ValueError(
            "unsupported table_rows keys: " + ", ".join(sorted(unsupported))
        )
    if "asset_liability" not in table_rows:
        return None
    rows = table_rows["asset_liability"]
    if not isinstance(rows, list):
        raise ValueError("table_rows.asset_liability must be an array")
    if not all(
        isinstance(row, list) and all(isinstance(cell, str) for cell in row)
        for row in rows
    ):
        raise ValueError(
            "table_rows.asset_liability must be a two-dimensional string array"
        )
    return rows


def _has_reusable_explicit_format(run) -> bool:
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None or run.font.size is None:
        return False
    return bool(r_pr.rFonts.get(qn("w:eastAsia")))


def _text_node_has_reusable_explicit_format(text_node) -> bool:
    run_element = text_node.getparent()
    if run_element is None or run_element.tag != qn("w:r"):
        return False
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        return False
    r_fonts = r_pr.find(qn("w:rFonts"))
    size = r_pr.find(qn("w:sz"))
    return bool(
        r_fonts is not None
        and r_fonts.get(qn("w:eastAsia"))
        and size is not None
        and size.get(qn("w:val"))
    )


def _prepare_empty_formatted_cells(table, rows: list[list[str]]) -> None:
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            if not str(value):
                continue
            cell = table.cell(row_index, col_index)
            text_nodes = list(cell._tc.iter(qn("w:t")))
            if text_nodes and _text_node_has_reusable_explicit_format(text_nodes[0]):
                continue
            reusable_run = next(
                (
                    run
                    for paragraph in cell.paragraphs
                    for run in paragraph.runs
                    if not run.text and _has_reusable_explicit_format(run)
                ),
                None,
            )
            if reusable_run is None:
                raise UpdateBlocked(
                    "no reusable formatted run; "
                    "first table text run has no reusable explicit format",
                    [f"row={row_index} column={col_index}"],
                )
            for text_node in text_nodes:
                parent = text_node.getparent()
                if parent is not None:
                    parent.remove(text_node)
            reusable_run._element.append(OxmlElement("w:t"))


def table_text_runs_have_explicit_format(
    table,
    rows: list[list[str]],
) -> bool:
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            if not str(value):
                continue
            cell = table.cell(row_index, col_index)
            first_text = next(iter(cell._tc.iter(qn("w:t"))), None)
            if (
                first_text is None
                or first_text.text != str(value)
                or not _text_node_has_reusable_explicit_format(first_text)
            ):
                return False
    return True


def _element_xml(element) -> str | None:
    return element.xml if element is not None else None


def table_format_fingerprint(table, rows: list[list[str]]) -> dict[str, object]:
    cells = []
    seen = set()
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            cells.append(
                {
                    "row": row_index,
                    "column": col_index,
                    "cell_properties": _element_xml(cell._tc.tcPr),
                    "paragraphs": [
                        {
                            "paragraph_properties": _element_xml(paragraph._p.pPr),
                            "run_properties": [
                                _element_xml(run._element.rPr)
                                for run in paragraph.runs
                            ],
                        }
                        for paragraph in cell.paragraphs
                    ],
                }
            )
    return {
        "table_properties": _element_xml(table._tbl.tblPr),
        "cells": cells,
    }


def planned_empty_cells_are_empty(table, rows: list[list[str]]) -> bool:
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            if str(value):
                continue
            cell = table.cell(row_index, col_index)
            if any((node.text or "") for node in cell._tc.iter(qn("w:t"))):
                return False
    return True


def update_financial_docx(
    source_docx: Path,
    bundle_path: Path,
    schema_path: Path,
    out_dir: Path,
) -> dict[str, str]:
    bundle = load_json(bundle_path)
    schema = load_json(schema_path)
    errors = validate_bundle(bundle, schema)
    if errors:
        raise UpdateBlocked("bundle validation failed", errors)
    asset_liability_rows = phase1_asset_liability_rows(bundle["docx_write_plan"])

    reject_out_dir_inside_skill_repo(out_dir)
    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    source_docx = source_docx.resolve()
    plan = bundle["docx_write_plan"]
    output = _safe_output_path(out_dir, plan["output_filename"], source_docx)
    planned_backup = backup_path(source_docx, out_dir, reserved=output)
    _ensure_distinct_document_paths(source_docx, output, planned_backup)
    backup = create_backup(source_docx, planned_backup)
    number_audit_path = out_dir / "number_audit.json"
    validation_result_path = out_dir / "validation_result.json"
    pending_output_path = out_dir / "待核验清单.md"
    write_json(number_audit_path, {"findings": []})
    staging: Path | None = None

    try:
        _ensure_distinct_document_paths(source_docx, output, backup)
        findings = audit_text(plan["analysis_markdown"], [bundle])
        if asset_liability_rows is not None:
            table_text = "\n".join(
                str(cell) for row in asset_liability_rows for cell in row
            )
            findings.extend(audit_text(table_text, [bundle]))
        number_audit = write_json(number_audit_path, {"findings": findings})
        if findings:
            raise UpdateBlocked("number audit failed", findings)

        document = Document(str(source_docx))
        locate_financial_section_anchors(
            document,
            plan["section_start"],
            plan["analysis_anchor"],
            plan["section_end"],
        )
        asset_table_index, table = find_table_in_section(
            document,
            ["资产负债简表", "资产总计"],
            plan["section_start"],
            plan["section_end"],
            preceding_anchor="合并数据",
        )
        expected_table_format = None
        if asset_liability_rows is not None:
            _validate_table_dimensions(table, asset_liability_rows)
            expected_table_format = table_format_fingerprint(
                table, asset_liability_rows
            )
            _prepare_empty_formatted_cells(table, asset_liability_rows)
            fill_table(table, asset_liability_rows)

        if plan["mode"] == "replace":
            location = replace_analysis_body(
                document=document,
                section_start=plan["section_start"],
                analysis_anchor=plan["analysis_anchor"],
                section_end=plan["section_end"],
                analysis_markdown=plan["analysis_markdown"],
                shading=plan["change_shading"],
            )
        else:
            location = insert_analysis_body(
                document=document,
                section_start=plan["section_start"],
                analysis_anchor=plan["analysis_anchor"],
                section_end=plan["section_end"],
                analysis_markdown=plan["analysis_markdown"],
                shading=plan["change_shading"],
            )
        _safe_output_path(out_dir, plan["output_filename"], source_docx)
        _ensure_distinct_document_paths(source_docx, output, backup)
        staging = _staging_output_path(output)
        document.save(str(staging))

        table_format_preserved = None
        table_text_runs_formatted = None
        planned_empty_cells_empty = None
        if asset_liability_rows is not None:
            saved_document = Document(str(staging))
            saved_table_index, saved_table = find_table_in_section(
                saved_document,
                ["资产负债简表", "资产总计"],
                plan["section_start"],
                plan["section_end"],
                preceding_anchor="合并数据",
            )
            table_format_preserved = (
                saved_table_index == asset_table_index
                and table_format_fingerprint(saved_table, asset_liability_rows)
                == expected_table_format
            )
            table_text_runs_formatted = table_text_runs_have_explicit_format(
                saved_table, asset_liability_rows
            )
            planned_empty_cells_empty = planned_empty_cells_are_empty(
                saved_table, asset_liability_rows
            )

        validation = validate_financial_docx(
            docx=staging,
            section_start=plan["section_start"],
            section_end=plan["section_end"],
            target_unit=plan["target_unit"],
            forbidden_units=["亿元"] if plan["target_unit"] == "万元" else [],
            expected_shading=plan["change_shading"],
            body_font="仿宋_GB2312",
            body_size=14,
            min_asset_table_rows=20,
            allow_codex_marker=plan["mode"] == "insert",
            analysis_anchor=plan["analysis_anchor"],
            expected_analysis_lines=plan["analysis_markdown"].splitlines(),
            asset_table_preceding_anchor="合并数据",
        )
        if table_format_preserved is not None:
            validation["checks"]["asset_table_target_format_preserved"] = (
                table_format_preserved
            )
            if not table_format_preserved:
                validation["failed"].append(
                    "asset_table_target_format_preserved"
                )
        if table_text_runs_formatted is not None:
            validation["checks"]["asset_table_target_text_runs_formatted"] = (
                table_text_runs_formatted
            )
            if not table_text_runs_formatted:
                validation["failed"].append(
                    "asset_table_target_text_runs_formatted"
                )
        if planned_empty_cells_empty is not None:
            validation["checks"]["asset_table_planned_empty_cells_empty"] = (
                planned_empty_cells_empty
            )
            if not planned_empty_cells_empty:
                validation["failed"].append(
                    "asset_table_planned_empty_cells_empty"
                )
        validation_result = write_json(validation_result_path, validation)
        pending_path = write_pending_markdown(
            pending_output_path, bundle["pending_verification"]
        )
        change_log = write_change_log(
            out=out_dir / "change_log.md",
            original=source_docx,
            workspace_copy=source_docx,
            backup=backup,
            output=output,
            mode=plan["mode"],
            location=location,
            sources=[source["name"] for source in bundle["sources"].values()],
            pending=[item["issue"] for item in bundle["pending_verification"]],
            table_preservation="original OOXML text-only update",
            validation_failed=list(validation["failed"]),
        )
        if validation["failed"]:
            raise UpdateFailed("DOCX validation failed", list(validation["failed"]))
        _safe_output_path(out_dir, plan["output_filename"], source_docx)
        result = {
            "backup": str(backup.resolve()),
            "output": str(output.resolve()),
            "number_audit": str(number_audit.resolve()),
            "validation_result": str(validation_result.resolve()),
            "change_log": str(change_log.resolve()),
            "pending_verification": str(pending_path.resolve()),
        }
        staging.replace(output)
        staging = None
        return result
    except Exception as error:
        if staging is not None:
            staging.unlink(missing_ok=True)
        if not validation_result_path.exists():
            write_json(
                validation_result_path,
                {"checks": {}, "failed": [str(error)]},
            )
        write_pending_markdown(
            pending_output_path,
            bundle["pending_verification"],
            runtime_failures=[str(error)],
        )
        raise


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source_docx", type=Path)
    parser.add_argument("bundle", type=Path)
    parser.add_argument("--schema", type=Path, required=True)
    parser.add_argument("--out-dir", type=Path, required=True)
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    try:
        result = update_financial_docx(
            source_docx=args.source_docx,
            bundle_path=args.bundle,
            schema_path=args.schema,
            out_dir=args.out_dir,
        )
    except (UpdateBlocked, UpdateFailed) as error:
        print(
            json.dumps(
                {"error": str(error), "details": error.details},
                ensure_ascii=False,
            ),
            file=sys.stderr,
        )
        return 2
    except (OSError, ValueError, json.JSONDecodeError) as error:
        print(
            json.dumps({"error": str(error)}, ensure_ascii=False),
            file=sys.stderr,
        )
        return 2
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
