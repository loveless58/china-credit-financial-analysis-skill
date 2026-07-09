#!/usr/bin/env python3
"""Extract paragraphs and financial-looking tables from a DOCX report."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


FINANCIAL_KEYWORDS = (
    "资产负债简表",
    "利润及利润分配表",
    "现金流量简表",
    "财务指标",
    "营业收入",
    "净利润",
    "资产总计",
    "负债合计",
    "所有者权益",
    "经营活动产生",
)


def cell_text(cell) -> str:
    return "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()).strip()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out-dir", type=Path, required=True)
    args = parser.parse_args()

    document = Document(str(args.docx))
    args.out_dir.mkdir(parents=True, exist_ok=True)

    paragraphs = [{"index": index, "text": paragraph.text} for index, paragraph in enumerate(document.paragraphs)]
    tables = []
    financial_tables = []
    for index, table in enumerate(document.tables):
        rows = [[cell_text(cell) for cell in row.cells] for row in table.rows]
        record = {"index": index, "rows": rows}
        tables.append(record)
        text = "\n".join("\t".join(row) for row in rows)
        if any(keyword in text for keyword in FINANCIAL_KEYWORDS):
            financial_tables.append(record)

    (args.out_dir / "paragraphs.json").write_text(json.dumps(paragraphs, ensure_ascii=False, indent=2), encoding="utf-8")
    (args.out_dir / "tables.json").write_text(json.dumps(tables, ensure_ascii=False, indent=2), encoding="utf-8")
    (args.out_dir / "financial_tables.json").write_text(
        json.dumps(financial_tables, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps({"paragraphs": len(paragraphs), "tables": len(tables), "financial_tables": len(financial_tables)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
