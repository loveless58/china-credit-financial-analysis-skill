#!/usr/bin/env python3
"""Validate a DOCX financial-analysis replacement structurally."""

from __future__ import annotations

import argparse
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def para_shade(paragraph) -> str | None:
    shd = paragraph._p.get_or_add_pPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def cell_shade(cell) -> str | None:
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def east_asia_font(run) -> str | None:
    r_fonts = run._element.get_or_add_rPr().rFonts
    return r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None


def table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def find_section(paragraphs: list[str], start_anchor: str, end_anchor: str) -> tuple[int, int]:
    start = next((i for i, text in enumerate(paragraphs) if start_anchor in text), None)
    if start is None:
        raise ValueError(f"section start not found: {start_anchor}")
    end = next((i for i, text in enumerate(paragraphs[start + 1 :], start + 1) if end_anchor in text), None)
    if end is None:
        raise ValueError(f"section end not found after start: {end_anchor}")
    return start, end


def first_run_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    return paragraph.runs[0]
    return None


def dominant_table_font(table) -> str | None:
    values = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = east_asia_font(run)
                    if font:
                        values.append(font)
    return Counter(values).most_common(1)[0][0] if values else None


def dominant_table_size(table) -> float | None:
    values = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is not None:
                        values.append(run.font.size.pt)
    return Counter(values).most_common(1)[0][0] if values else None


def find_expected_analysis_paragraphs(
    document,
    start: int,
    end: int,
    analysis_anchor: str,
    expected_lines: list[str],
) -> tuple[bool, list]:
    anchor_index = next(
        (
            index
            for index in range(start, end)
            if analysis_anchor in document.paragraphs[index].text
        ),
        None,
    )
    if anchor_index is None:
        return False, []

    matched = []
    cursor = anchor_index + 1
    for expected_line in expected_lines:
        match_index = next(
            (
                index
                for index in range(cursor, end)
                if document.paragraphs[index].text.strip() == expected_line
            ),
            None,
        )
        if match_index is None:
            break
        matched.append(document.paragraphs[match_index])
        cursor = match_index + 1
    return True, matched


def paragraph_font_matches(paragraph, expected_font: str) -> bool:
    text_runs = [run for run in paragraph.runs if run.text.strip()]
    return bool(text_runs) and all(
        east_asia_font(run) == expected_font for run in text_runs
    )


def paragraph_size_matches(paragraph, expected_size: float) -> bool:
    text_runs = [run for run in paragraph.runs if run.text.strip()]
    return bool(text_runs) and all(
        run.font.size is not None
        and abs(run.font.size.pt - expected_size) < 0.001
        for run in text_runs
    )


def validate_financial_docx(
    docx: Path,
    section_start: str,
    section_end: str,
    target_unit: str,
    forbidden_units: list[str],
    expected_shading: str,
    body_font: str | None = None,
    body_size: float | None = None,
    table_font: str | None = None,
    table_size: float | None = None,
    min_asset_table_rows: int = 20,
    allow_codex_marker: bool = False,
    analysis_anchor: str | None = None,
    expected_analysis_lines: list[str] | None = None,
) -> dict[str, object]:
    docx = docx.resolve()
    if not docx.is_file() or docx.suffix.lower() != ".docx":
        return {"checks": {"docx_exists": False}, "failed": ["docx_exists"]}

    doc = Document(str(docx))
    para_texts = [paragraph.text for paragraph in doc.paragraphs]
    start, end = find_section(para_texts, section_start, section_end)
    section_text = "\n".join(para_texts[start:end])
    financial_tables = [
        table
        for table in doc.tables
        if any(
            anchor in table_text(table)
            for anchor in ("资产负债简表", "资产总计", "财务指标")
        )
    ]
    asset_tables = [
        table
        for table in financial_tables
        if "资产负债简表" in table_text(table) or "资产总计" in table_text(table)
    ]
    indicator_tables = [
        table for table in financial_tables if "财务指标" in table_text(table)
    ]
    shaded_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs[start:end]
        if para_shade(paragraph) == expected_shading
    ]
    precise_body_validation = (
        analysis_anchor is not None and expected_analysis_lines is not None
    )
    expected_lines = [
        line.strip()
        for line in (expected_analysis_lines or [])
        if line.strip()
    ]
    anchor_found = False
    matched_analysis_paragraphs = []
    if precise_body_validation and expected_lines:
        anchor_found, matched_analysis_paragraphs = find_expected_analysis_paragraphs(
            doc,
            start,
            end,
            analysis_anchor,
            expected_lines,
        )
    body_sample_paragraphs = (
        matched_analysis_paragraphs if precise_body_validation else shaded_paragraphs
    )
    body_run = next(
        (
            run
            for paragraph in body_sample_paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ),
        None,
    )
    table_run = first_run_in_table(asset_tables[0]) if asset_tables else None
    detected_table_font = dominant_table_font(asset_tables[0]) if asset_tables else None
    detected_table_size = dominant_table_size(asset_tables[0]) if asset_tables else None

    checks: dict[str, object] = {
        "docx_exists": True,
        "financial_section_found": start < end,
        "target_unit_present": target_unit in section_text
        or any(target_unit in table_text(table) for table in financial_tables),
        "forbidden_unit_absent_in_section": not any(
            unit in section_text for unit in forbidden_units
        ),
        "asset_liability_table_found": bool(asset_tables),
        "asset_liability_table_rows": len(asset_tables[0].rows) if asset_tables else 0,
        "asset_liability_table_cols": len(asset_tables[0].columns) if asset_tables else 0,
        "asset_liability_table_not_too_simple": bool(asset_tables)
        and len(asset_tables[0].rows) >= min_asset_table_rows,
        "indicator_table_found": bool(indicator_tables),
        "codex_marker_absent": "Codex" not in section_text,
        "section_shaded_paragraphs": len(shaded_paragraphs),
        "section_shading_present": bool(shaded_paragraphs),
        "asset_table_first_cell_shading": cell_shade(asset_tables[0].cell(0, 0))
        if asset_tables
        else None,
        "body_font_sample": east_asia_font(body_run) if body_run else None,
        "body_size_sample": body_run.font.size.pt
        if body_run is not None and body_run.font.size is not None
        else None,
        "table_font_sample": detected_table_font
        or (east_asia_font(table_run) if table_run else None),
        "table_size_sample": detected_table_size
        or (
            table_run.font.size.pt
            if table_run is not None and table_run.font.size is not None
            else None
        ),
    }
    if analysis_anchor is not None or expected_analysis_lines is not None:
        all_lines_found = (
            precise_body_validation
            and bool(expected_lines)
            and len(matched_analysis_paragraphs) == len(expected_lines)
        )
        checks["analysis_anchor_found"] = anchor_found
        checks["analysis_lines_found"] = all_lines_found
        checks["analysis_paragraph_shading_matches"] = all_lines_found and all(
            para_shade(paragraph) == expected_shading
            for paragraph in matched_analysis_paragraphs
        )
        checks["analysis_paragraph_font_matches"] = (
            all_lines_found
            and body_font is not None
            and all(
                paragraph_font_matches(paragraph, body_font)
                for paragraph in matched_analysis_paragraphs
            )
        )
        checks["analysis_paragraph_size_matches"] = (
            all_lines_found
            and body_size is not None
            and all(
                paragraph_size_matches(paragraph, body_size)
                for paragraph in matched_analysis_paragraphs
            )
        )
    if body_font is not None:
        checks["body_font_matches"] = checks["body_font_sample"] == body_font
    if body_size is not None:
        checks["body_size_matches"] = checks["body_size_sample"] == body_size
    if table_font is not None:
        checks["table_font_matches"] = checks["table_font_sample"] == table_font
    if table_size is not None:
        checks["table_size_matches"] = checks["table_size_sample"] == table_size

    required = [
        "docx_exists",
        "financial_section_found",
        "target_unit_present",
        "forbidden_unit_absent_in_section",
        "asset_liability_table_found",
        "asset_liability_table_not_too_simple",
        "section_shading_present",
    ]
    if not allow_codex_marker:
        required.append("codex_marker_absent")
    if body_font is not None:
        required.append("body_font_matches")
    if body_size is not None:
        required.append("body_size_matches")
    if table_font is not None:
        required.append("table_font_matches")
    if table_size is not None:
        required.append("table_size_matches")
    if analysis_anchor is not None or expected_analysis_lines is not None:
        required.extend(
            [
                "analysis_anchor_found",
                "analysis_lines_found",
                "analysis_paragraph_shading_matches",
                "analysis_paragraph_font_matches",
                "analysis_paragraph_size_matches",
            ]
        )
    failed = [key for key in required if not checks.get(key)]
    return {"checks": checks, "failed": failed}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--section-start", default="财务分析")
    parser.add_argument("--section-end", default="行业分析")
    parser.add_argument("--target-unit", default="万元")
    parser.add_argument("--forbidden-unit", action="append", default=["亿元"])
    parser.add_argument("--expected-shading", default="FFF2CC")
    parser.add_argument("--body-font")
    parser.add_argument("--body-size", type=float)
    parser.add_argument("--table-font")
    parser.add_argument("--table-size", type=float)
    parser.add_argument("--min-asset-table-rows", type=int, default=20)
    parser.add_argument("--allow-codex-marker", action="store_true")
    parser.add_argument("--analysis-anchor")
    parser.add_argument("--expected-analysis-line", action="append")
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()

    try:
        result = validate_financial_docx(
            docx=args.docx,
            section_start=args.section_start,
            section_end=args.section_end,
            target_unit=args.target_unit,
            forbidden_units=args.forbidden_unit,
            expected_shading=args.expected_shading,
            body_font=args.body_font,
            body_size=args.body_size,
            table_font=args.table_font,
            table_size=args.table_size,
            min_asset_table_rows=args.min_asset_table_rows,
            allow_codex_marker=args.allow_codex_marker,
            analysis_anchor=args.analysis_anchor,
            expected_analysis_lines=args.expected_analysis_line,
        )
    except ValueError as error:
        result = {"checks": {}, "failed": [str(error)]}
    text = json.dumps(result, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(text, encoding="utf-8")
    print(text)
    return 1 if result["failed"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
