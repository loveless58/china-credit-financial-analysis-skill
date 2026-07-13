#!/usr/bin/env python3
"""End-to-end self-test for the safe financial DOCX update workflow."""

from __future__ import annotations

import copy
import hashlib
import json
import os
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from backup_docx import backup_path
import update_financial_docx as updater_module
from validate_financial_docx import validate_financial_docx


ROOT = Path(__file__).resolve().parent.parent
UPDATER = ROOT / "scripts" / "update_financial_docx.py"
SCHEMA = ROOT / "schemas" / "financial_analysis_bundle.schema.json"
SECTION_START = "（五）、财务分析"
ANALYSIS_ANCHOR = "合并财务情况分析："
SECTION_END = "（六）、行业分析"
OUTPUT_FILENAME = "某测试公司_财务分析更新稿.docx"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file_handle:
        for chunk in iter(lambda: file_handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "00AA00")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def east_asia_font(run) -> str | None:
    r_fonts = run._element.get_or_add_rPr().rFonts
    return r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None


def cell_fill(cell) -> str | None:
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def paragraph_fill(paragraph) -> str | None:
    shd = paragraph._p.get_or_add_pPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_run(run, font_name: str, font_size: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), font_name
    )


def add_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.add_run(text)
    return inserted


def table_rows() -> list[list[str]]:
    rows = [
        ["资产负债简表（合并口径，单位：万元）", "", ""],
        ["项目", "2024", "2025"],
        ["资产总计", "100.00", "120.00"],
    ]
    rows.extend([["待核验项目", "", ""] for _ in range(17)])
    return rows


def make_source_docx(path: Path, target_cell_mode: str = "text") -> None:
    document = Document()
    document.add_paragraph(SECTION_START)
    document.add_paragraph("合并数据：")
    table = document.add_table(rows=20, cols=3)
    table.style = "Table Grid"
    set_borders(table)
    source_rows = table_rows()
    source_rows[2][2] = "110.00"
    for row_index, row in enumerate(source_rows):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_shading(cell, "D9EAD3")
            if row_index == 2 and col_index == 2 and target_cell_mode != "text":
                if target_cell_mode == "formatted-empty-run":
                    format_run(cell.paragraphs[0].add_run(), "仿宋", 11)
                elif target_cell_mode in {
                    "unsafe-empty-text-node",
                    "unsafe-text-node-with-formatted-run",
                }:
                    unsafe_run = cell.paragraphs[0].add_run()
                    unsafe_run._element.append(OxmlElement("w:t"))
                    if target_cell_mode == "unsafe-text-node-with-formatted-run":
                        format_run(cell.paragraphs[0].add_run(), "仿宋", 11)
                elif target_cell_mode != "unsafe-empty":
                    raise ValueError(f"unknown target_cell_mode: {target_cell_mode}")
                continue
            format_run(cell.paragraphs[0].add_run(text), "仿宋", 11)
    document.add_paragraph(ANALYSIS_ANCHOR)
    document.add_paragraph("旧财务分析内容。")
    document.add_paragraph(SECTION_END)
    document.save(path)


def valid_bundle() -> dict:
    return {
        "schema_version": "1.0",
        "company_name": "某测试公司",
        "reporting_basis": "合并口径",
        "currency": "CNY",
        "unit": "万元",
        "periods": ["2024", "2025"],
        "sources": {
            "annual_2025": {
                "name": "某测试公司2025年审计报告",
                "type": "audit_report",
                "file": "source/annual_2025.pdf",
            }
        },
        "financial_tables": {
            "asset_liability": {
                "rows": [
                    {
                        "metric": "total_assets",
                        "label": "资产总计",
                        "values": {
                            "2024": {
                                "value": "100.00",
                                "status": "verified",
                                "source_refs": ["annual_2025"],
                            },
                            "2025": {
                                "value": "120.00",
                                "status": "verified",
                                "source_refs": ["annual_2025"],
                            },
                        },
                    }
                ]
            }
        },
        "ratios": {},
        "risk_points": [
            {
                "category": "偿债能力",
                "statement": "需关注债务期限结构与经营现金流匹配情况。",
                "evidence_refs": ["annual_2025"],
            }
        ],
        "pending_verification": [
            {"issue": "核验受限资产明细。", "status": "source_missing"}
        ],
        "docx_write_plan": {
            "mode": "replace",
            "section_start": SECTION_START,
            "analysis_anchor": ANALYSIS_ANCHOR,
            "section_end": SECTION_END,
            "analysis_markdown": "总资产为120.00万元。",
            "table_rows": {"asset_liability": table_rows()},
            "target_unit": "万元",
            "require_backup": True,
            "preserve_asset_liability_table": True,
            "change_shading": "FFF2CC",
            "output_filename": OUTPUT_FILENAME,
        },
    }


def run_updater(source: Path, bundle: dict, out_dir: Path) -> subprocess.CompletedProcess[str]:
    bundle_path = out_dir.parent / f"{out_dir.name}.bundle.json"
    bundle_path.write_text(json.dumps(bundle, ensure_ascii=False), encoding="utf-8")
    environment = os.environ.copy()
    environment["PYTHONIOENCODING"] = "utf-8"
    return subprocess.run(
        [
            sys.executable,
            "-B",
            str(UPDATER),
            str(source),
            str(bundle_path),
            "--schema",
            str(SCHEMA),
            "--out-dir",
            str(out_dir),
        ],
        cwd=ROOT,
        env=environment,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )


def assert_replace_success(source: Path, source_hash_before: str, tmp_path: Path) -> None:
    out_dir = tmp_path / "replace-output"
    result = run_updater(source, valid_bundle(), out_dir)
    assert result.returncode == 0, result.stderr or result.stdout
    summary = json.loads(result.stdout)

    updated_docx = out_dir / OUTPUT_FILENAME
    assert Path(summary["output"]) == updated_docx.resolve()
    assert sha256(source) == source_hash_before
    assert len(list(out_dir.glob("*.backup-*.docx"))) == 1
    assert updated_docx.exists()

    output_document = Document(str(updated_docx))
    output_text = "\n".join(paragraph.text for paragraph in output_document.paragraphs)
    assert "旧财务分析内容" not in output_text
    assert "总资产为120.00万元" in output_text
    output_table = output_document.tables[0]
    assert output_table.cell(2, 2).text == "120.00"
    output_run = output_table.cell(2, 2).paragraphs[0].runs[0]
    assert east_asia_font(output_run) == "仿宋"
    assert output_run.font.size.pt == 11
    assert cell_fill(output_table.cell(2, 2)) == "D9EAD3"
    assert "00AA00" in output_table._tbl.tblPr.xml
    assert json.loads((out_dir / "number_audit.json").read_text(encoding="utf-8"))["findings"] == []
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert validation["failed"] == []
    for key in (
        "analysis_anchor_found",
        "analysis_lines_found",
        "analysis_paragraph_shading_matches",
        "analysis_paragraph_font_matches",
        "analysis_paragraph_size_matches",
    ):
        assert validation["checks"][key] is True
    assert validation["checks"]["asset_table_target_format_preserved"] is True
    assert validation["checks"]["asset_table_target_text_runs_formatted"] is True
    assert validation["checks"]["asset_table_planned_empty_cells_empty"] is True
    assert validation["checks"]["asset_table_first_cell_shading"] == "D9EAD3"
    assert (out_dir / "change_log.md").exists()
    assert (out_dir / "待核验清单.md").exists()


def assert_number_audit_failure(source: Path, tmp_path: Path) -> None:
    out_dir = tmp_path / "audit-failure-output"
    bundle = valid_bundle()
    bundle["docx_write_plan"]["analysis_markdown"] = "总资产为999.00万元。"
    result = run_updater(source, bundle, out_dir)
    assert result.returncode != 0
    audit = json.loads((out_dir / "number_audit.json").read_text(encoding="utf-8"))
    assert [finding["number"] for finding in audit["findings"]] == ["999.00"]
    assert len(list(out_dir.glob("*.backup-*.docx"))) == 1
    assert not (out_dir / OUTPUT_FILENAME).exists()
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert "number audit failed" in validation["failed"]
    pending = (out_dir / "待核验清单.md").read_text(encoding="utf-8")
    assert "核验受限资产明细" in pending
    assert "number audit failed" in pending


def assert_insert_success(source: Path, tmp_path: Path) -> None:
    out_dir = tmp_path / "insert-output"
    bundle = copy.deepcopy(valid_bundle())
    bundle["docx_write_plan"]["mode"] = "insert"
    result = run_updater(source, bundle, out_dir)
    assert result.returncode == 0, result.stderr or result.stdout

    output_document = Document(str(out_dir / OUTPUT_FILENAME))
    paragraphs = output_document.paragraphs
    texts = [paragraph.text for paragraph in paragraphs]
    anchor_index = texts.index(ANALYSIS_ANCHOR)
    assert texts[anchor_index + 1] == "【Codex财务分析更新建议】"
    assert texts[anchor_index + 2] == "总资产为120.00万元。"
    assert "旧财务分析内容。" in texts
    assert paragraph_fill(paragraphs[anchor_index + 1]) == "FFF2CC"
    assert paragraph_fill(paragraphs[anchor_index + 2]) == "FFF2CC"
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert validation["failed"] == []
    assert validation["checks"]["codex_marker_absent"] is False
    for key in (
        "analysis_anchor_found",
        "analysis_lines_found",
        "analysis_paragraph_shading_matches",
        "analysis_paragraph_font_matches",
        "analysis_paragraph_size_matches",
    ):
        assert validation["checks"][key] is True


def assert_output_entity_safety(tmp_path: Path) -> None:
    hardlink_source = tmp_path / "hardlink-source.docx"
    make_source_docx(hardlink_source)
    source_hash_before = sha256(hardlink_source)
    hardlink_out = tmp_path / "hardlink-output"
    hardlink_out.mkdir()
    output = hardlink_out / OUTPUT_FILENAME
    os.link(hardlink_source, output)

    result = run_updater(hardlink_source, valid_bundle(), hardlink_out)
    assert result.returncode != 0, "hard-linked output was accepted"
    assert "same file entity as source" in result.stderr
    assert sha256(hardlink_source) == source_hash_before
    assert not list(hardlink_out.glob("*.backup-*.docx"))
    assert not (hardlink_out / "number_audit.json").exists()

    existing_source = tmp_path / "existing-source.docx"
    make_source_docx(existing_source)
    existing_hash_before = sha256(existing_source)
    existing_out = tmp_path / "existing-output"
    existing_out.mkdir()
    occupied = existing_out / OUTPUT_FILENAME
    occupied.write_bytes(b"occupied")

    result = run_updater(existing_source, valid_bundle(), existing_out)
    assert result.returncode != 0, "existing output was accepted"
    assert "output path already exists" in result.stderr
    assert occupied.read_bytes() == b"occupied"
    assert sha256(existing_source) == existing_hash_before
    assert not list(existing_out.glob("*.backup-*.docx"))
    assert not (existing_out / "number_audit.json").exists()


def assert_backup_output_conflict_protection(tmp_path: Path) -> None:
    source = tmp_path / "backup-conflict-source.docx"
    make_source_docx(source)
    out_dir = tmp_path / "backup-conflict-output"
    out_dir.mkdir()
    reserved_output = backup_path(source, out_dir)
    guarded_backup = backup_path(source, out_dir, reserved=reserved_output)
    assert guarded_backup != reserved_output
    assert guarded_backup.parent == reserved_output.parent
    assert guarded_backup.suffix == ".docx"


def assert_table_dimension_failures(tmp_path: Path) -> None:
    source = tmp_path / "dimension-source.docx"
    make_source_docx(source)
    source_hash_before = sha256(source)
    base_rows = table_rows()
    cases = {
        "fewer-rows": base_rows[:-1],
        "extra-rows": [*base_rows, ["待核验项目", "", ""]],
        "short-row": [*base_rows[:3], ["待核验项目", ""], *base_rows[4:]],
        "long-row": [
            *base_rows[:3],
            ["待核验项目", "", "", ""],
            *base_rows[4:],
        ],
    }
    for name, rows in cases.items():
        bundle = valid_bundle()
        bundle["docx_write_plan"]["table_rows"]["asset_liability"] = rows
        out_dir = tmp_path / f"dimension-{name}"
        result = run_updater(source, bundle, out_dir)
        assert result.returncode != 0, f"invalid table dimensions accepted: {name}"
        assert "table dimensions do not match template" in result.stderr
        assert len(list(out_dir.glob("*.backup-*.docx"))) == 1
        assert (out_dir / "number_audit.json").exists()
        assert not (out_dir / OUTPUT_FILENAME).exists()
        assert sha256(source) == source_hash_before


def run_with_text(cell, text: str):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if text in run.text:
                return run
    raise AssertionError(f"run with text not found: {text}")


def assert_formatted_empty_run_preserved(tmp_path: Path) -> None:
    source = tmp_path / "formatted-empty-source.docx"
    make_source_docx(source, target_cell_mode="formatted-empty-run")
    out_dir = tmp_path / "formatted-empty-output"
    result = run_updater(source, valid_bundle(), out_dir)
    assert result.returncode == 0, result.stderr or result.stdout

    output_document = Document(str(out_dir / OUTPUT_FILENAME))
    output_table = output_document.tables[0]
    output_run = run_with_text(output_table.cell(2, 2), "120.00")
    assert east_asia_font(output_run) == "仿宋"
    assert output_run.font.size.pt == 11
    assert cell_fill(output_table.cell(2, 2)) == "D9EAD3"
    assert "00AA00" in output_table._tbl.tblPr.xml
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert validation["checks"]["asset_table_target_format_preserved"] is True
    assert validation["checks"]["asset_table_target_text_runs_formatted"] is True

    unsafe_source = tmp_path / "unsafe-empty-source.docx"
    make_source_docx(unsafe_source, target_cell_mode="unsafe-empty")
    unsafe_out = tmp_path / "unsafe-empty-output"
    unsafe_result = run_updater(unsafe_source, valid_bundle(), unsafe_out)
    assert unsafe_result.returncode != 0, "unformatted empty cell was accepted"
    assert "no reusable formatted run" in unsafe_result.stderr
    assert len(list(unsafe_out.glob("*.backup-*.docx"))) == 1
    assert (unsafe_out / "number_audit.json").exists()
    assert not (unsafe_out / OUTPUT_FILENAME).exists()


def assert_unrelated_shading_not_accepted(tmp_path: Path) -> None:
    docx = tmp_path / "unrelated-shading.docx"
    make_source_docx(docx)
    document = Document(str(docx))
    heading = document.paragraphs[0]
    set_paragraph_shading(heading, "FFF2CC")
    format_run(heading.runs[0], "仿宋_GB2312", 14)
    document.save(str(docx))

    validation = validate_financial_docx(
        docx=docx,
        section_start=SECTION_START,
        section_end=SECTION_END,
        target_unit="万元",
        forbidden_units=["亿元"],
        expected_shading="FFF2CC",
        body_font="仿宋_GB2312",
        body_size=14,
        min_asset_table_rows=20,
        analysis_anchor=ANALYSIS_ANCHOR,
        expected_analysis_lines=["旧财务分析内容。"],
    )
    assert "analysis_paragraph_shading_matches" in validation["failed"]
    assert "analysis_paragraph_font_matches" in validation["failed"]
    assert "analysis_paragraph_size_matches" in validation["failed"]


def assert_unsafe_empty_text_node_handled(tmp_path: Path) -> None:
    unsafe_source = tmp_path / "unsafe-text-node-source.docx"
    make_source_docx(unsafe_source, target_cell_mode="unsafe-empty-text-node")
    unsafe_out = tmp_path / "unsafe-text-node-output"
    unsafe_result = run_updater(unsafe_source, valid_bundle(), unsafe_out)
    assert unsafe_result.returncode != 0, "unsafe empty w:t was accepted"
    assert "first table text run has no reusable explicit format" in unsafe_result.stderr
    assert len(list(unsafe_out.glob("*.backup-*.docx"))) == 1
    assert (unsafe_out / "number_audit.json").exists()
    assert not (unsafe_out / OUTPUT_FILENAME).exists()

    fallback_source = tmp_path / "safe-fallback-source.docx"
    make_source_docx(
        fallback_source,
        target_cell_mode="unsafe-text-node-with-formatted-run",
    )
    fallback_out = tmp_path / "safe-fallback-output"
    fallback_result = run_updater(fallback_source, valid_bundle(), fallback_out)
    assert fallback_result.returncode == 0, fallback_result.stderr
    output_document = Document(str(fallback_out / OUTPUT_FILENAME))
    output_run = run_with_text(output_document.tables[0].cell(2, 2), "120.00")
    assert east_asia_font(output_run) == "仿宋"
    assert output_run.font.size.pt == 11
    validation = json.loads(
        (fallback_out / "validation_result.json").read_text(encoding="utf-8")
    )
    assert validation["checks"]["asset_table_target_text_runs_formatted"] is True


def assert_noncontiguous_analysis_rejected(tmp_path: Path) -> None:
    docx = tmp_path / "noncontiguous-analysis.docx"
    make_source_docx(docx)
    document = Document(str(docx))
    wrong_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "旧财务分析内容。"
    )
    wrong_paragraph.text = "错误中间段落。"
    expected_paragraph = add_paragraph_after(wrong_paragraph, "旧财务分析内容。")
    set_paragraph_shading(expected_paragraph, "FFF2CC")
    format_run(expected_paragraph.runs[0], "仿宋_GB2312", 14)
    document.save(str(docx))

    validation = validate_financial_docx(
        docx=docx,
        section_start=SECTION_START,
        section_end=SECTION_END,
        target_unit="万元",
        forbidden_units=["亿元"],
        expected_shading="FFF2CC",
        body_font="仿宋_GB2312",
        body_size=14,
        min_asset_table_rows=20,
        allow_codex_marker=False,
        analysis_anchor=ANALYSIS_ANCHOR,
        expected_analysis_lines=["旧财务分析内容。"],
    )
    assert validation["checks"]["analysis_lines_found"] is False
    assert "analysis_lines_found" in validation["failed"]


def assert_phase1_plan_guard() -> None:
    plan = valid_bundle()["docx_write_plan"]
    phase1_asset_liability_rows = updater_module.phase1_asset_liability_rows
    assert phase1_asset_liability_rows(plan) == table_rows()

    no_table_update = copy.deepcopy(plan)
    no_table_update["table_rows"] = {}
    assert phase1_asset_liability_rows(no_table_update) is None

    unsupported = copy.deepcopy(plan)
    unsupported["table_rows"]["profit"] = [["利润表"]]
    try:
        phase1_asset_liability_rows(unsupported)
    except ValueError as error:
        assert "unsupported table_rows keys" in str(error)
    else:
        raise AssertionError("unsupported Phase 1 table_rows key was accepted")

    preservation_disabled = copy.deepcopy(plan)
    preservation_disabled["preserve_asset_liability_table"] = False
    try:
        phase1_asset_liability_rows(preservation_disabled)
    except ValueError as error:
        assert "preserve_asset_liability_table must equal true" in str(error)
    else:
        raise AssertionError("disabled asset-liability preservation was accepted")


def assert_empty_table_rows_success(tmp_path: Path) -> None:
    source = tmp_path / "empty-table-rows-source.docx"
    make_source_docx(source)
    bundle = valid_bundle()
    bundle["docx_write_plan"]["table_rows"] = {}
    out_dir = tmp_path / "empty-table-rows-output"
    result = run_updater(source, bundle, out_dir)
    assert result.returncode == 0, result.stderr
    output = Document(str(out_dir / OUTPUT_FILENAME))
    assert output.tables[0].cell(2, 2).text == "110.00"
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert validation["failed"] == []
    assert "asset_table_target_format_preserved" not in validation["checks"]


def add_decoy_table_before_section(path: Path) -> None:
    document = Document(str(path))
    decoy_xml = deepcopy(document.tables[0]._tbl)
    first_text = next(iter(decoy_xml.iter(qn("w:t"))))
    first_text.text = "资产负债简表（无关章节，单位：万元）"
    document._body._body.insert(0, decoy_xml)
    document.save(str(path))


def assert_scoped_table_selection(tmp_path: Path) -> None:
    source = tmp_path / "scoped-table-source.docx"
    make_source_docx(source)
    add_decoy_table_before_section(source)
    out_dir = tmp_path / "scoped-table-output"
    result = run_updater(source, valid_bundle(), out_dir)
    assert result.returncode == 0, result.stderr

    output = Document(str(out_dir / OUTPUT_FILENAME))
    assert output.tables[0].cell(2, 2).text == "110.00"
    assert output.tables[1].cell(2, 2).text == "120.00"

    multiple_source = tmp_path / "multiple-scoped-tables.docx"
    make_source_docx(multiple_source)
    multiple_doc = Document(str(multiple_source))
    multiple_doc.tables[0]._tbl.addnext(deepcopy(multiple_doc.tables[0]._tbl))
    multiple_doc.save(str(multiple_source))
    multiple_out = tmp_path / "multiple-scoped-output"
    multiple_result = run_updater(multiple_source, valid_bundle(), multiple_out)
    assert multiple_result.returncode != 0
    assert "found 2" in multiple_result.stderr
    assert len(list(multiple_out.glob("*.backup-*.docx"))) == 1
    assert not (multiple_out / OUTPUT_FILENAME).exists()
    multiple_pending = (multiple_out / "待核验清单.md").read_text(encoding="utf-8")
    assert "found 2" in multiple_pending
    assert "核验受限资产明细" in multiple_pending

    missing_source = tmp_path / "missing-scoped-table.docx"
    make_source_docx(missing_source)
    add_decoy_table_before_section(missing_source)
    missing_doc = Document(str(missing_source))
    missing_doc.tables[1].cell(0, 0).text = "其他报表"
    missing_doc.tables[1].cell(2, 0).text = "其他项目"
    missing_doc.save(str(missing_source))
    missing_out = tmp_path / "missing-scoped-output"
    missing_result = run_updater(missing_source, valid_bundle(), missing_out)
    assert missing_result.returncode != 0
    assert "found 0" in missing_result.stderr
    assert len(list(missing_out.glob("*.backup-*.docx"))) == 1
    assert not (missing_out / OUTPUT_FILENAME).exists()
    missing_pending = (missing_out / "待核验清单.md").read_text(encoding="utf-8")
    assert "found 0" in missing_pending
    assert "核验受限资产明细" in missing_pending


def assert_scoped_table_unit_validation(tmp_path: Path) -> None:
    forbidden_docx = tmp_path / "forbidden-table-unit.docx"
    make_source_docx(forbidden_docx)
    forbidden_doc = Document(str(forbidden_docx))
    forbidden_doc.tables[0].cell(0, 0).text = "资产负债简表（合并口径，单位：亿元）"
    forbidden_doc.save(str(forbidden_docx))
    forbidden_validation = validate_financial_docx(
        docx=forbidden_docx,
        section_start=SECTION_START,
        section_end=SECTION_END,
        target_unit="万元",
        forbidden_units=["亿元"],
        expected_shading="FFF2CC",
        min_asset_table_rows=20,
    )
    assert forbidden_validation["checks"]["forbidden_unit_absent_in_section"] is False

    unrelated_docx = tmp_path / "unrelated-target-unit.docx"
    make_source_docx(unrelated_docx)
    unrelated_doc = Document(str(unrelated_docx))
    unrelated_doc.tables[0].cell(0, 0).text = "资产负债简表（合并口径）"
    unrelated_doc.paragraphs[-2].text = "正文单位为万元。"
    unrelated_doc.save(str(unrelated_docx))
    add_decoy_table_before_section(unrelated_docx)
    unrelated_validation = validate_financial_docx(
        docx=unrelated_docx,
        section_start=SECTION_START,
        section_end=SECTION_END,
        target_unit="万元",
        forbidden_units=["亿元"],
        expected_shading="FFF2CC",
        min_asset_table_rows=20,
    )
    assert unrelated_validation["checks"]["target_unit_present"] is True
    assert (
        unrelated_validation["checks"]["asset_table_title_contains_target_unit"]
        is False
    )
    assert "asset_table_title_contains_target_unit" in unrelated_validation["failed"]


def assert_pending_written_after_anchor_failure(tmp_path: Path) -> None:
    source = tmp_path / "anchor-failure-source.docx"
    make_source_docx(source)
    bundle = valid_bundle()
    bundle["docx_write_plan"]["analysis_anchor"] = "不存在的正文锚点"
    out_dir = tmp_path / "anchor-failure-output"
    result = run_updater(source, bundle, out_dir)
    assert result.returncode != 0
    assert len(list(out_dir.glob("*.backup-*.docx"))) == 1
    assert not (out_dir / OUTPUT_FILENAME).exists()
    pending = (out_dir / "待核验清单.md").read_text(encoding="utf-8")
    assert "核验受限资产明细" in pending
    assert "不存在的正文锚点" in pending or "analysis anchor" in pending


def assert_body_anchor_boundaries(tmp_path: Path) -> None:
    cross_section_source = tmp_path / "cross-section-anchor-source.docx"
    make_source_docx(cross_section_source)
    cross_section_doc = Document(str(cross_section_source))
    early_end = cross_section_doc.add_paragraph(SECTION_END)
    cross_section_doc.tables[0]._tbl.addnext(early_end._p)
    cross_section_doc.save(str(cross_section_source))

    direct_bundle = tmp_path / "cross-section-anchor.bundle.json"
    direct_bundle.write_text(
        json.dumps(valid_bundle(), ensure_ascii=False), encoding="utf-8"
    )
    original_fill_table = updater_module.fill_table

    def reject_early_table_mutation(*_args, **_kwargs):
        raise AssertionError("table mutation started before body anchor validation")

    updater_module.fill_table = reject_early_table_mutation
    try:
        try:
            updater_module.update_financial_docx(
                source_docx=cross_section_source,
                bundle_path=direct_bundle,
                schema_path=SCHEMA,
                out_dir=tmp_path / "cross-section-direct-output",
            )
        except ValueError as error:
            assert "analysis anchor" in str(error)
        else:
            raise AssertionError("cross-section analysis anchor was accepted")
    finally:
        updater_module.fill_table = original_fill_table

    cross_section_out = tmp_path / "cross-section-anchor-output"
    cross_section_result = run_updater(
        cross_section_source, valid_bundle(), cross_section_out
    )
    assert cross_section_result.returncode != 0
    assert not (cross_section_out / OUTPUT_FILENAME).exists()

    duplicate_source = tmp_path / "duplicate-analysis-anchor-source.docx"
    make_source_docx(duplicate_source)
    duplicate_doc = Document(str(duplicate_source))
    add_paragraph_after(duplicate_doc.paragraphs[-2], ANALYSIS_ANCHOR)
    duplicate_doc.save(str(duplicate_source))
    duplicate_out = tmp_path / "duplicate-analysis-anchor-output"
    duplicate_result = run_updater(duplicate_source, valid_bundle(), duplicate_out)
    assert duplicate_result.returncode != 0
    assert not (duplicate_out / OUTPUT_FILENAME).exists()

    try:
        validate_financial_docx(
            docx=duplicate_source,
            section_start=SECTION_START,
            section_end=SECTION_END,
            target_unit="万元",
            forbidden_units=["亿元"],
            expected_shading="FFF2CC",
            min_asset_table_rows=20,
            analysis_anchor=ANALYSIS_ANCHOR,
            expected_analysis_lines=["旧财务分析内容。"],
            asset_table_preceding_anchor="合并数据",
        )
    except ValueError as error:
        assert "analysis anchor" in str(error)
    else:
        raise AssertionError("validator accepted duplicate analysis anchors")


def assert_staging_cleanup_and_atomic_publish(tmp_path: Path) -> None:
    source = tmp_path / "staging-validation-source.docx"
    make_source_docx(source)
    source_document = Document(str(source))
    source_document.paragraphs[-2].text = "旧正文金额单位为亿元。"
    source_document.save(str(source))
    source_hash_before = sha256(source)

    bundle = valid_bundle()
    bundle["docx_write_plan"]["mode"] = "insert"
    out_dir = tmp_path / "staging-validation-output"
    result = run_updater(source, bundle, out_dir)
    assert result.returncode != 0
    assert sha256(source) == source_hash_before
    assert len(list(out_dir.glob("*.backup-*.docx"))) == 1
    assert (out_dir / "number_audit.json").exists()
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert "forbidden_unit_absent_in_section" in validation["failed"]
    assert (out_dir / "待核验清单.md").exists()
    assert not (out_dir / OUTPUT_FILENAME).exists()
    assert not [path for path in out_dir.iterdir() if ".staging-" in path.name]


def assert_fixed_artifact_conflict_blocks_before_backup(tmp_path: Path) -> None:
    source = tmp_path / "artifact-conflict-source.docx"
    make_source_docx(source)
    out_dir = tmp_path / "artifact-conflict-output"
    out_dir.mkdir()
    (out_dir / "number_audit.json").mkdir()

    result = run_updater(source, valid_bundle(), out_dir)
    assert result.returncode != 0
    assert not list(out_dir.glob("*.backup-*.docx"))
    assert not (out_dir / OUTPUT_FILENAME).exists()
    assert not [path for path in out_dir.iterdir() if ".staging-" in path.name]


def assert_table_format_fingerprint_detects_loss(tmp_path: Path) -> None:
    from update_financial_docx import table_format_fingerprint

    docx = tmp_path / "table-fingerprint.docx"
    make_source_docx(docx)
    document = Document(str(docx))
    table = document.tables[0]
    expected = table_format_fingerprint(table, table_rows())
    set_cell_shading(table.cell(3, 1), "ABCDEF")
    changed = table_format_fingerprint(table, table_rows())
    assert changed != expected

    row_document = Document(str(docx))
    row_table = row_document.tables[0]
    row_expected = table_format_fingerprint(row_table, table_rows())
    row_table.rows[0].height = Pt(42)
    assert table_format_fingerprint(row_table, table_rows()) != row_expected

    grid_document = Document(str(docx))
    grid_table = grid_document.tables[0]
    grid_expected = table_format_fingerprint(grid_table, table_rows())
    grid_col = grid_table._tbl.tblGrid.find(qn("w:gridCol"))
    assert grid_col is not None
    grid_col.set(qn("w:w"), "9999")
    assert table_format_fingerprint(grid_table, table_rows()) != grid_expected


def assert_planned_empty_cells_are_validated(tmp_path: Path) -> None:
    source = tmp_path / "planned-empty-source.docx"
    make_source_docx(source)
    out_dir = tmp_path / "planned-empty-output"
    result = run_updater(source, valid_bundle(), out_dir)
    assert result.returncode == 0, result.stderr
    output = Document(str(out_dir / OUTPUT_FILENAME))
    assert output.tables[0].cell(3, 1).text == ""
    validation = json.loads(
        (out_dir / "validation_result.json").read_text(encoding="utf-8")
    )
    assert validation["checks"]["asset_table_planned_empty_cells_empty"] is True


def run_selected_case(case: str, tmp_path: Path) -> None:
    if case == "hardlink-output":
        assert_output_entity_safety(tmp_path)
    elif case == "backup-output-conflict":
        assert_backup_output_conflict_protection(tmp_path)
    elif case == "table-dimensions":
        assert_table_dimension_failures(tmp_path)
    elif case == "formatted-empty-run":
        assert_formatted_empty_run_preserved(tmp_path)
    elif case == "unrelated-shading":
        assert_unrelated_shading_not_accepted(tmp_path)
    elif case == "table-format-gate":
        assert_table_format_fingerprint_detects_loss(tmp_path)
    elif case == "planned-empty-cells":
        assert_planned_empty_cells_are_validated(tmp_path)
    elif case == "unsafe-empty-text-node":
        assert_unsafe_empty_text_node_handled(tmp_path)
    elif case == "noncontiguous-analysis":
        assert_noncontiguous_analysis_rejected(tmp_path)
    elif case == "phase1-plan-guard":
        assert_phase1_plan_guard()
    elif case == "empty-table-rows":
        assert_empty_table_rows_success(tmp_path)
    elif case == "scoped-table-selection":
        assert_scoped_table_selection(tmp_path)
    elif case == "scoped-table-units":
        assert_scoped_table_unit_validation(tmp_path)
    elif case == "pending-after-failure":
        assert_pending_written_after_anchor_failure(tmp_path)
    elif case == "body-anchor-boundaries":
        assert_body_anchor_boundaries(tmp_path)
    elif case == "staging-atomic-publish":
        assert_staging_cleanup_and_atomic_publish(tmp_path)
    elif case == "artifact-path-conflict":
        assert_fixed_artifact_conflict_blocks_before_backup(tmp_path)
    else:
        raise ValueError(f"unknown self-test case: {case}")


def main() -> int:
    selected_case = sys.argv[1] if len(sys.argv) > 1 else "all"
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        if selected_case != "all":
            run_selected_case(selected_case, tmp_path)
            print(f"financial DOCX update self-test case passed: {selected_case}")
            return 0

        source = tmp_path / "source.docx"
        make_source_docx(source)
        source_hash_before = sha256(source)

        assert_replace_success(source, source_hash_before, tmp_path)
        assert_number_audit_failure(source, tmp_path)
        assert_insert_success(source, tmp_path)
        assert sha256(source) == source_hash_before
        assert_output_entity_safety(tmp_path)
        assert_backup_output_conflict_protection(tmp_path)
        assert_table_dimension_failures(tmp_path)
        assert_formatted_empty_run_preserved(tmp_path)
        assert_unrelated_shading_not_accepted(tmp_path)
        assert_table_format_fingerprint_detects_loss(tmp_path)
        assert_planned_empty_cells_are_validated(tmp_path)
        assert_unsafe_empty_text_node_handled(tmp_path)
        assert_noncontiguous_analysis_rejected(tmp_path)
        assert_phase1_plan_guard()
        assert_empty_table_rows_success(tmp_path)
        assert_scoped_table_selection(tmp_path)
        assert_scoped_table_unit_validation(tmp_path)
        assert_pending_written_after_anchor_failure(tmp_path)
        assert_body_anchor_boundaries(tmp_path)
        assert_staging_cleanup_and_atomic_publish(tmp_path)
        assert_fixed_artifact_conflict_blocks_before_backup(tmp_path)

    print("financial DOCX update self-test passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
