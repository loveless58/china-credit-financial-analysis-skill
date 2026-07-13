#!/usr/bin/env python3
"""Self-test for preserving DOCX asset-liability table formatting."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from preserve_financial_table_format import find_table_in_section


SCRIPT = Path(__file__).with_name("preserve_financial_table_format.py")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "00AA00")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def east_asia_font(run) -> str | None:
    r_fonts = run._element.get_or_add_rPr().rFonts
    return r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None


def run_with_text(cell, text: str):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if text in run.text:
                return run
    raise AssertionError(f"run with text not found: {text}")


def cell_fill(cell) -> str | None:
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def add_asset_table(doc, amount: str = "200"):
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    set_borders(table)
    rows = [
        ["资产负债简表（合并口径，单位：万元）", "", ""],
        ["项目", "2024.12", "2025.12"],
        ["资产总计", "100", amount],
    ]
    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            set_cell_shading(cell, "D9EAD3")
            run = cell.paragraphs[0].add_run(text)
            run.font.name = "仿宋"
            run.font.size = Pt(11)
            run.bold = row_index <= 1
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "仿宋")
    return table


def make_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("（五）、财务分析")
    add_asset_table(doc)
    doc.add_paragraph("（六）、行业分析")
    doc.save(path)


def assert_section_scoped_locator(tmp_path: Path) -> None:
    path = tmp_path / "section-scoped.docx"
    doc = Document()
    decoy = add_asset_table(doc, "999")
    doc.add_paragraph("（五）、财务分析")
    target = add_asset_table(doc, "200")
    doc.add_paragraph("（六）、行业分析")
    doc.save(path)

    loaded = Document(str(path))
    table_index, found = find_table_in_section(
        loaded,
        ["资产负债简表", "资产总计"],
        "（五）、财务分析",
        "（六）、行业分析",
    )
    assert table_index == 1
    assert "200" in "\n".join(c.text for r in found.rows for c in r.cells)
    assert "999" in "\n".join(c.text for r in loaded.tables[0].rows for c in r.cells)

    no_match = Document()
    no_match.add_paragraph("（五）、财务分析")
    no_match.add_table(rows=1, cols=1).cell(0, 0).text = "其他表"
    no_match.add_paragraph("（六）、行业分析")
    try:
        find_table_in_section(
            no_match,
            ["资产负债简表", "资产总计"],
            "（五）、财务分析",
            "（六）、行业分析",
        )
    except ValueError as error:
        assert "found 0" in str(error)
    else:
        raise AssertionError("zero scoped candidates were accepted")

    multiple = Document()
    multiple.add_paragraph("（五）、财务分析")
    add_asset_table(multiple, "200")
    add_asset_table(multiple, "300")
    multiple.add_paragraph("（六）、行业分析")
    try:
        find_table_in_section(
            multiple,
            ["资产负债简表", "资产总计"],
            "（五）、财务分析",
            "（六）、行业分析",
        )
    except ValueError as error:
        assert "found 2" in str(error)
    else:
        raise AssertionError("multiple scoped candidates were accepted")

    contextual = Document()
    contextual.add_paragraph("（五）、财务分析")
    contextual.add_paragraph("  合并数据 ：  ")
    expected_contextual = add_asset_table(contextual, "200")
    contextual.add_paragraph("本部财务数据：")
    add_asset_table(contextual, "300")
    contextual.add_paragraph("（六）、行业分析")
    contextual_index, contextual_target = find_table_in_section(
        contextual,
        ["资产负债简表", "资产总计"],
        "（五）、财务分析",
        "（六）、行业分析",
        preceding_anchor="合并数据",
    )
    assert contextual_index == 0
    assert contextual_target._tbl is expected_contextual._tbl

    parent_only = Document()
    parent_only.add_paragraph("（五）、财务分析")
    parent_only.add_paragraph("本部财务数据：")
    add_asset_table(parent_only, "300")
    parent_only.add_paragraph("（六）、行业分析")
    try:
        find_table_in_section(
            parent_only,
            ["资产负债简表", "资产总计"],
            "（五）、财务分析",
            "（六）、行业分析",
            preceding_anchor="合并数据",
        )
    except ValueError as error:
        assert "found 0" in str(error)
    else:
        raise AssertionError("single parent-company table bypassed preceding_anchor")


def main() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        template = tmp_path / "template.docx"
        output = tmp_path / "output.docx"
        rows_json = tmp_path / "rows.json"
        make_template(template)
        rows_json.write_text(
            json.dumps(
                {
                    "asset_liability": [
                        ["资产负债简表（合并口径，单位：万元）", "", ""],
                        ["项目", "2024.12", "2025.12"],
                        ["资产总计", "370737", "375593"],
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        subprocess.run(
            [
                sys.executable,
                str(SCRIPT),
                str(template),
                str(rows_json),
                "--out",
                str(output),
                "--table-key",
                "asset_liability",
                "--anchor",
                "资产负债简表",
            ],
            check=True,
        )
        doc = Document(str(output))
        source_doc = Document(str(template))
        source_table = source_doc.tables[0]
        target = next(t for t in doc.tables if "375593" in "\n".join(c.text for r in t.rows for c in r.cells))
        source_run = run_with_text(source_table.cell(2, 2), "200")
        target_run = run_with_text(target.cell(2, 2), "375593")
        assert east_asia_font(target_run) == east_asia_font(source_run) == "仿宋"
        assert target_run.font.size.pt == source_run.font.size.pt == 11
        assert cell_fill(target.cell(2, 2)) == cell_fill(source_table.cell(2, 2)) == "D9EAD3"
        assert target.cell(2, 2).text == "375593"
        assert_section_scoped_locator(tmp_path)
        print("table format preservation self-test passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
